import io
import re
import base64
from typing import Dict, List, Tuple
from pathlib import Path

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    print("[DocProcessor] PyPDF2 not available")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[DocProcessor] python-docx not available")


def detect_encoded_payloads(text: str) -> List[Dict]:
    """
    Detect encoded payloads (Base64, Hex, ROT13, etc).
    
    Returns:
        List of detected encoded patterns
    """
    detections = []
    
    # Base64 detection (40+ characters)
    base64_pattern = r'[A-Za-z0-9+/]{40,}={0,2}'
    base64_matches = re.finditer(base64_pattern, text)
    for match in base64_matches:
        try:
            decoded = base64.b64decode(match.group()).decode('utf-8', errors='ignore')
            if len(decoded) > 10:
                detections.append({
                    "type": "base64",
                    "encoded": match.group()[:50],
                    "decoded_preview": decoded[:100]
                })
        except:
            pass
    
    # Hex encoding detection
    hex_pattern = r'0x[0-9a-fA-F]{20,}'
    hex_matches = re.finditer(hex_pattern, text)
    for match in hex_matches:
        detections.append({
            "type": "hex",
            "encoded": match.group()[:50]
        })
    
    # ROT13 detection (common obfuscation)
    # Check for suspicious words encoded in ROT13
    suspicious_keywords = [
        "ignore previous", "system prompt", "bypass", "override",
        "unrestricted", "dan mode", "jailbreak", "reveal hidden"
    ]
    
    for keyword in suspicious_keywords:
        # Simple ROT13 check
        rot13_keyword = ''.join(
            chr((ord(c) - ord('a') + 13) % 26 + ord('a')) if c.islower()
            else chr((ord(c) - ord('A') + 13) % 26 + ord('A')) if c.isupper()
            else c for c in keyword
        )
        if rot13_keyword.lower() in text.lower():
            detections.append({
                "type": "rot13_suspicious",
                "pattern": rot13_keyword
            })
    
    return detections


def extract_urls(text: str) -> List[str]:
    """Extract URLs from text."""
    url_pattern = r'https?://[^\s\)\]\}\'"<>]+'
    return re.findall(url_pattern, text)


def detect_jailbreak_attempts(text: str) -> List[str]:
    """Detect common jailbreak and prompt injection attempts."""
    jailbreak_patterns = [
        r"ignore\s+previous\s+instructions?",
        r"disregard\s+(all\s+)?previous",
        r"act\s+as\s+dan",
        r"unrestricted\s+mode",
        r"developer\s+mode",
        r"system\s*:\s*reveal",
        r"hidden\s+rules",
        r"bypass\s+(all\s+)?filters?",
        r"override\s+(all\s+)?restrictions?",
        r"forget\s+all\s+previous",
        r"reset\s+(the\s+)?system",
        r"you\s+are\s+now",
        r"pretend\s+you\s+are",
        r"roleplay\s+as"
    ]
    
    detections = []
    for pattern in jailbreak_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            detections.append(pattern)
    
    return detections


def detect_macro_indicators(text: str) -> List[str]:
    """Detect macro and script indicators."""
    macro_patterns = [
        r"Sub\s+\w+\(",
        r"Function\s+\w+\(",
        r"Private\s+Sub",
        r"Public\s+Sub",
        r"<script[^>]*>",
        r"javascript:",
        r"VBScript",
        r"ActiveXObject",
        r"WScript\.",
        r"ShellExecute"
    ]
    
    detections = []
    for pattern in macro_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            detections.append(pattern)
    
    return detections


def extract_pdf(file_bytes: bytes) -> Tuple[str, Dict]:
    """
    Extract text from PDF file.
    
    Args:
        file_bytes: Raw PDF data
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    if not PDF_AVAILABLE:
        return "", {"error": "PyPDF2 not available"}
    
    try:
        print("[DocProcessor] Extracting text from PDF...")
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_file)
        
        metadata = {
            "pages": len(reader.pages),
            "has_metadata": bool(reader.metadata),
            "annotations": 0,
            "urls_found": [],
            "suspicious_indicators": []
        }
        
        # Extract metadata if present
        if reader.metadata:
            metadata["title"] = reader.metadata.get("/Title", "")
            metadata["author"] = reader.metadata.get("/Author", "")
            metadata["subject"] = reader.metadata.get("/Subject", "")
            metadata["creator"] = reader.metadata.get("/Creator", "")
        
        extracted_text = []
        
        # Extract text from each page
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                extracted_text.append(page_text)
                
                # Check for annotations
                if "/Annots" in page:
                    metadata["annotations"] += len(page["/Annots"])
            except Exception as e:
                print(f"[DocProcessor] Error extracting page {page_num}: {e}")
        
        full_text = "\n".join(extracted_text)
        
        # Extract URLs
        urls = extract_urls(full_text)
        if urls:
            metadata["urls_found"] = urls
        
        # Check for embedded JavaScript (PDF XSS)
        if "/JavaScript" in str(reader):
            metadata["suspicious_indicators"].append("PDF contains embedded JavaScript")
        
        print(f"[DocProcessor] Extracted {len(full_text)} characters from {metadata['pages']} pages")
        
        return full_text, metadata
    except Exception as e:
        print(f"[DocProcessor] Error processing PDF: {e}")
        return "", {"error": str(e)}


def extract_docx(file_bytes: bytes) -> Tuple[str, Dict]:
    """
    Extract text from DOCX file.
    
    Args:
        file_bytes: Raw DOCX data
        
    Returns:
        Tuple of (extracted_text, metadata_dict)
    """
    if not DOCX_AVAILABLE:
        return "", {"error": "python-docx not available"}
    
    try:
        print("[DocProcessor] Extracting text from DOCX...")
        
        docx_file = io.BytesIO(file_bytes)
        doc = DocxDocument(docx_file)
        
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
            "core_properties": {},
            "comments": 0,
            "macros": False,
            "suspicious_indicators": []
        }
        
        # Extract core properties
        if doc.core_properties:
            metadata["core_properties"] = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or ""
            }
        
        extracted_text = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                extracted_text.append(para.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        extracted_text.append(cell.text)
        
        # Check for comments
        if hasattr(doc, 'comments_part') and doc.comments_part:
            metadata["comments"] = len(doc.comments_part.comments)
        
        # Check for macros (DOCM files)
        if file_bytes.startswith(b'PK'):  # ZIP format
            try:
                import zipfile
                with zipfile.ZipFile(io.BytesIO(file_bytes)) as z:
                    if 'word/vbaProject.bin' in z.namelist():
                        metadata["macros"] = True
                        metadata["suspicious_indicators"].append("Document contains VBA macros")
            except:
                pass
        
        full_text = "\n".join(extracted_text)
        
        print(f"[DocProcessor] Extracted {len(full_text)} characters from DOCX")
        
        return full_text, metadata
    except Exception as e:
        print(f"[DocProcessor] Error processing DOCX: {e}")
        return "", {"error": str(e)}


def extract_text_file(file_bytes: bytes, encoding: str = 'utf-8') -> Tuple[str, Dict]:
    """Extract text from plain text file."""
    try:
        text = file_bytes.decode(encoding)
        return text, {"encoding": encoding, "length": len(text)}
    except:
        try:
            text = file_bytes.decode('latin-1')
            return text, {"encoding": "latin-1", "length": len(text)}
        except:
            return "", {"error": "Could not decode file"}


def analyze_document(file_bytes: bytes, filename: str) -> Dict:
    """
    Comprehensive document analysis.
    
    Args:
        file_bytes: Raw file data
        filename: Original filename
        
    Returns:
        dict with complete analysis results
    """
    try:
        file_ext = Path(filename).suffix.lower()
        
        print(f"[DocProcessor] Analyzing document: {filename} ({file_ext})")
        
        # Determine file type and extract
        if file_ext == '.pdf':
            extracted_text, metadata = extract_pdf(file_bytes)
        elif file_ext == '.docx':
            extracted_text, metadata = extract_docx(file_bytes)
        elif file_ext in ['.txt', '.rtf']:
            extracted_text, metadata = extract_text_file(file_bytes)
        elif file_ext == '.doc':
            # For legacy DOC files, treat as text
            extracted_text, metadata = extract_text_file(file_bytes)
        else:
            return {
                "success": False,
                "error": f"Unsupported file type: {file_ext}",
                "extracted_text": ""
            }
        
        if not extracted_text:
            return {
                "success": False,
                "error": metadata.get("error", "Could not extract text from document"),
                "extracted_text": ""
            }
        
        # Analyze extracted text for threats
        encoded_payloads = detect_encoded_payloads(extracted_text)
        jailbreak_attempts = detect_jailbreak_attempts(extracted_text)
        macro_indicators = detect_macro_indicators(extracted_text)
        urls = extract_urls(extracted_text)
        
        # Calculate threat score
        threat_score = 0
        threat_indicators = []
        
        if jailbreak_attempts:
            threat_score += 30
            threat_indicators.extend(jailbreak_attempts)
        
        if encoded_payloads:
            threat_score += 20
            threat_indicators.append(f"{len(encoded_payloads)} encoded payloads detected")
        
        if macro_indicators:
            threat_score += 40
            threat_indicators.extend(macro_indicators)
        
        if metadata.get("suspicious_indicators"):
            threat_score += 15
            threat_indicators.extend(metadata["suspicious_indicators"])
        
        if urls:
            threat_score += len(urls) * 5
            threat_indicators.append(f"{len(urls)} URLs found")
        
        threat_score = min(100, threat_score)
        
        result = {
            "success": True,
            "extracted_text": extracted_text,
            "file_type": file_ext,
            "metadata": metadata,
            "threat_analysis": {
                "threat_score": threat_score,
                "jailbreak_attempts": jailbreak_attempts,
                "encoded_payloads": encoded_payloads,
                "macro_indicators": macro_indicators,
                "urls_found": urls,
                "threat_indicators": threat_indicators
            },
            "analysis_summary": {
                "total_characters": len(extracted_text),
                "detected_threats": len(threat_indicators),
                "encoded_payloads": len(encoded_payloads),
                "jailbreak_attempts": len(jailbreak_attempts),
                "macros_found": len(macro_indicators) > 0
            }
        }
        
        print(f"[DocProcessor] Analysis complete - Threat: {threat_score}/100")
        
        return result
    except Exception as e:
        print(f"[DocProcessor] Error in document analysis: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "extracted_text": ""
        }
